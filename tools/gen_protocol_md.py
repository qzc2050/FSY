#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""从辐射报警仪协议 docx/xlsx 生成 Markdown 文档。"""

import os
import re
import sys
from pathlib import Path

from docx import Document
from openpyxl import load_workbook

ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = ROOT / "docs" / "protocol"

FILES = {
    "spec": ROOT / "辐射报警仪协议说明09.docx",
    "commands": ROOT / "辐射报警仪协议命令和寄存器表09.xlsx",
    "debug": ROOT / "辐射报警仪协议调试示例09.docx",
}


def para_to_md(text: str) -> str:
    text = text.replace("\xa0", " ").strip()
    if not text:
        return ""
    # 标题样式推断
    if re.match(r"^[一二三四五六七八九十]+[\s　]", text):
        return f"## {text.lstrip('　 ')}"
    if re.match(r"^\d+\s", text):
        return f"### {text.lstrip('　 ')}"
    return text


def docx_to_md(path: Path) -> str:
    doc = Document(str(path))
    lines: list[str] = [f"# {path.stem}", "", f"> 源文件：`{path.name}`", ""]

    for p in doc.paragraphs:
        md = para_to_md(p.text)
        if md:
            lines.append(md)
            lines.append("")

    for ti, table in enumerate(doc.tables, 1):
        rows = []
        for row in table.rows:
            cells = [c.text.replace("\n", " ").replace("|", "\\|").strip() for c in row.cells]
            if any(cells):
                rows.append(cells)
        if not rows:
            continue
        lines.append(f"### 表格 {ti}")
        lines.append("")
        header = rows[0]
        lines.append("| " + " | ".join(header) + " |")
        lines.append("| " + " | ".join("---" for _ in header) + " |")
        for row in rows[1:]:
            while len(row) < len(header):
                row.append("")
            lines.append("| " + " | ".join(row[: len(header)]) + " |")
        lines.append("")

    return "\n".join(lines).rstrip() + "\n"


def sheet_to_md_table(ws, max_cols: int | None = None) -> str:
    rows: list[list[str]] = []
    for row in ws.iter_rows(values_only=True):
        cells = ["" if c is None else str(c).replace("\n", " ").replace("|", "\\|").strip() for c in row]
        if max_cols:
            cells = cells[:max_cols]
        # 去掉尾部空列
        while cells and not cells[-1]:
            cells.pop()
        if any(cells):
            rows.append(cells)

    if not rows:
        return "_（空表）_\n"

    col_count = max(len(r) for r in rows)
    lines: list[str] = []
    header = rows[0] + [""] * (col_count - len(rows[0]))
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join("---" for _ in range(col_count)) + " |")
    for row in rows[1:]:
        padded = row + [""] * (col_count - len(row))
        lines.append("| " + " | ".join(padded) + " |")
    return "\n".join(lines) + "\n"


def xlsx_to_md(path: Path) -> str:
    wb = load_workbook(str(path), data_only=True)
    lines: list[str] = [
        f"# {path.stem}",
        "",
        f"> 源文件：`{path.name}`",
        "",
    ]

    sheet_titles = {
        "命令表": "命令与帧格式",
        "寄存器表": "寄存器定义",
        "协议": "协议概要",
        "Sheet1": "其他",
    }

    for sn in wb.sheetnames:
        ws = wb[sn]
        title = sheet_titles.get(sn, sn)
        lines.append(f"## {title}（{sn}）")
        lines.append("")
        # 命令表列很多，保留前 14 列主要字段
        max_cols = 14 if sn == "命令表" else None
        lines.append(sheet_to_md_table(ws, max_cols=max_cols))
        lines.append("")

    return "\n".join(lines).rstrip() + "\n"


def debug_docx_to_md(path: Path) -> str:
    doc = Document(str(path))
    lines: list[str] = [
        f"# {path.stem}",
        "",
        f"> 源文件：`{path.name}`",
        "",
    ]

    current_section = ""
    for p in doc.paragraphs:
        text = p.text.replace("\xa0", " ").strip()
        if not text:
            continue

        if re.match(r"^[一二三四五六七八九十]+", text):
            current_section = text
            lines.append(f"## {text}")
            lines.append("")
            continue

        if re.match(r"^\d+\s", text):
            lines.append(f"### {text}")
            lines.append("")
            continue

        # 十六进制帧
        if re.search(r"\b[0-9A-Fa-f]{2}(?:\s+[0-9A-Fa-f]{2}){7,}\b", text):
            hex_part = re.search(r"([0-9A-Fa-f]{2}(?:\s+[0-9A-Fa-f]{2})+)", text).group(1)
            desc = text.replace(hex_part, "").strip()
            lines.append("```")
            lines.append(hex_part.upper())
            lines.append("```")
            if desc:
                lines.append("")
                lines.append(desc)
            lines.append("")
            continue

        lines.append(text)
        lines.append("")

    for ti, table in enumerate(doc.tables, 1):
        rows = []
        for row in table.rows:
            cells = [c.text.replace("\n", " ").strip() for c in row.cells]
            if any(cells):
                rows.append(cells)
        if not rows:
            continue
        lines.append(f"## 表格 {ti}")
        lines.append("")
        header = rows[0]
        lines.append("| " + " | ".join(header) + " |")
        lines.append("| " + " | ".join("---" for _ in header) + " |")
        for row in rows[1:]:
            while len(row) < len(header):
                row.append("")
            lines.append("| " + " | ".join(row[: len(header)]) + " |")
        lines.append("")

    return "\n".join(lines).rstrip() + "\n"


def main() -> int:
    sys.stdout.reconfigure(encoding="utf-8")
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    outputs = {
        "辐射报警仪协议说明09.md": docx_to_md(FILES["spec"]),
        "辐射报警仪协议命令和寄存器表09.md": xlsx_to_md(FILES["commands"]),
        "辐射报警仪协议调试示例09.md": debug_docx_to_md(FILES["debug"]),
    }

    for name, content in outputs.items():
        out_path = OUT_DIR / name
        out_path.write_text(content, encoding="utf-8")
        print(f"Wrote {out_path} ({len(content)} chars)")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
